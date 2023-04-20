from docx import Document
from docx.shared import Inches
import os

'''
Programmer: JB Andreski
Date Made: 4/18/23
How to Use: Put as many .png files in the folder that this program is in as you want.
            You need to have a file name 'Blank.png' that is a blank png file for this to work
            Run the program and .docx file will appear in the folder
'''
#Gets the absolute path to the folder the file is in
cwd = os.getcwd() 

#Gets all the files necesary for the program from the folder
files = []
for file in os.listdir(cwd):
    if file.endswith(".png") and file != "Blank.png":
        files.append(file)
    elif file == "Blank.png":
        blank = file

#Gets what the user wants to file to be named
docName = input("Enter What You Want the File to Be Named: ")

#Caluclates how many pages the document will have
numPages = int(len(files)/4)
if (len(files)%4) != 0:
     numPages += 1

document = Document() #Makes the document

#Sets the margins, Change these numbers for grid alignment
for section in document.sections:
        section.top_margin = Inches(0.25)
        section.bottom_margin = Inches(0.25)
        section.left_margin = Inches(0.25)
        section.right_margin = Inches(0.25)


a = 0
b = 1
c = 2
d = 3
#Loop through each page that needs to be made
for x in range(numPages):
    if files[b] != "Blank.png":
        paragraph = document.add_paragraph(files[a] +  "\t\t" + files[b]) 
        paragraph.alignment = 1
    else:
        paragraph = document.add_paragraph("\t\t" +files[a])
        paragraph.alignment = 0
    table = document.add_table(rows = 10, cols = 8) #Makes the table
    rows = -1
    cols = 0
    #Fills the table
    for row in table.rows:
       cols = 0
       rows += 1
       for cell in row.cells:
            #Each if statement is a location check to ensure the right image goes into the right place
            if cols == 3   or cols == 4 or rows == 4 or rows == 5:
                paragraph = cell.paragraphs[0]
                run = paragraph.add_run()
                run.add_picture(os.path.join(cwd, "Blank.png"), width=Inches(0.75), height=Inches(0.75)) #Change the width and height for spacing
            elif cols < 3 and rows < 4:
                paragraph = cell.paragraphs[0]
                run = paragraph.add_run()
                run.add_picture(os.path.join(cwd, files[a]), width=Inches(0.75), height=Inches(0.75))
            elif cols > 4 and rows < 4:
                paragraph = cell.paragraphs[0]
                run = paragraph.add_run()
                run.add_picture(os.path.join(cwd, files[b]), width=Inches(0.75), height=Inches(0.75))
            elif cols < 3 and rows > 5:
                paragraph = cell.paragraphs[0]
                run = paragraph.add_run()
                run.add_picture(os.path.join(cwd, files[c]), width=Inches(0.75), height=Inches(0.75))
            elif cols > 4 and rows > 5:
                paragraph = cell.paragraphs[0]
                run = paragraph.add_run()
                run.add_picture(os.path.join(cwd, files[d]), width=Inches(0.75), height=Inches(0.75))
            cols += 1
    if files[d] != "Blank.png":
        paragraph = document.add_paragraph(files[c] +  "\t\t" + files[d])
        paragraph.alignment = 1
    elif files[c] != "Blank.png":
        paragraph = document.add_paragraph("\t\t" + files[c])
        paragraph.alignment = 0
    document.add_page_break()
    a += 4
    b += 4
    c += 4
    d += 4
    #Adds blank pictures to the end of the files array to fill the rest of the grid with blanks
    if b > (len(files) -1):
        files.append(blank)
        files.append(blank)
        files.append(blank)
    elif c > (len(files) -1):
        files.append(blank)
        files.append(blank)
    elif d > (len(files) -1):
        files.append(blank)




document.save(docName + ".docx") #Saves the document with the file name given
